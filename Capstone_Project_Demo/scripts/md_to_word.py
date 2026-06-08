"""Convert every .md file in D:/Capstone_Project_Demo (recursively) to a styled
Microsoft Word document and save them in D:/Capstone_Project_Demo/presentation_file/
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── paths ─────────────────────────────────────────────────────────────────────
ROOT_DIR   = Path(r"D:\Capstone_Project_Demo")
OUTPUT_DIR = ROOT_DIR / "presentation_file"
OUTPUT_DIR.mkdir(exist_ok=True)

# Directories to exclude from the search
EXCLUDE_DIRS = {".venv", "node_modules", ".git", "__pycache__", "chroma_db"}

# ── colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x39, 0x64)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
TEAL   = RGBColor(0x17, 0x75, 0x6A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x26, 0x26, 0x26)
GREY   = RGBColor(0x59, 0x59, 0x59)

# ── docx helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F3964")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


def _new_doc(title: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(2.0))

    style = doc.styles["Normal"]
    style.font.name      = "Calibri"
    style.font.size      = Pt(10.5)
    style.font.color.rgb = DARK

    # Cover header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run(title)
    r.bold           = True
    r.font.size      = Pt(22)
    r.font.color.rgb = NAVY

    _add_hr(doc)
    return doc


def _h1(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(15)
    r.font.color.rgb = NAVY
    _add_hr(doc)


def _h2(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(13)
    r.font.color.rgb = BLUE


def _h3(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(11.5)
    r.font.color.rgb = TEAL


def _body(doc, text: str, bold=False, italic=False, code=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if code:
        p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    if code:
        r.font.name      = "Courier New"
        r.font.size      = Pt(9)
        r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        rPr = r._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "EEF2F7")
        rPr.append(shd)
    else:
        r.font.size      = Pt(10.5)
        r.font.color.rgb = DARK


def _bullet(doc, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2 + level * 0.2)
    r = p.add_run(text)
    r.font.size      = Pt(10.5)
    r.font.color.rgb = DARK


def _inline_styled(doc, raw: str):
    """
    Add a body paragraph with inline bold (**text**) and inline code (`text`) rendering.
    Falls back to plain if neither marker found.
    """
    # Split on **bold** or `code`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", raw)
    if len(parts) == 1:
        _body(doc, raw)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold           = True
            r.font.size      = Pt(10.5)
            r.font.color.rgb = DARK
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name      = "Courier New"
            r.font.size      = Pt(9.5)
            r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        else:
            r = p.add_run(part)
            r.font.size      = Pt(10.5)
            r.font.color.rgb = DARK


def _table(doc, header_row: list[str], data_rows: list[list[str]]):
    n     = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=n)
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0].cells
    for i, h in enumerate(header_row):
        _set_cell_bg(hdr[i], "1F3964")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r   = p.add_run(h.strip())
        r.bold           = True
        r.font.size      = Pt(9.5)
        r.font.color.rgb = WHITE

    for ri, row_data in enumerate(data_rows):
        fill  = "D6E4F0" if ri % 2 == 0 else "FFFFFF"
        cells = table.rows[ri + 1].cells
        for ci in range(n):
            txt = row_data[ci].strip() if ci < len(row_data) else ""
            _set_cell_bg(cells[ci], fill)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            r = cells[ci].paragraphs[0].add_run(txt)
            r.font.size      = Pt(9.5)
            r.font.color.rgb = DARK

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


# ── markdown parser → docx ────────────────────────────────────────────────────

def _strip_inline(text: str) -> str:
    """Strip common markdown inline markers for plain text rendering."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text.strip()


def md_to_docx(md_path: Path, out_path: Path):
    raw   = md_path.read_text(encoding="utf-8", errors="replace")
    lines = raw.splitlines()

    # Derive document title from first H1 or filename
    title = md_path.stem.replace("_", " ")
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            break

    doc = _new_doc(title)

    # ── table accumulator ─────────────────────────────────────────────────────
    in_table      = False
    table_headers: list[str]       = []
    table_rows:    list[list[str]] = []

    # ── code block accumulator ────────────────────────────────────────────────
    in_code  = False
    code_buf: list[str] = []

    def flush_table():
        nonlocal in_table, table_headers, table_rows
        if table_headers and table_rows:
            _table(doc, table_headers, table_rows)
        in_table      = False
        table_headers = []
        table_rows    = []

    def flush_code():
        nonlocal in_code, code_buf
        for cl in code_buf:
            _body(doc, cl if cl else " ", code=True)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(3)
        in_code  = False
        code_buf = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── fenced code block ─────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                if in_table:
                    flush_table()
                in_code  = True
                code_buf = []
            else:
                flush_code()
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── markdown table row ────────────────────────────────────────────────
        if line.strip().startswith("|"):
            cols = [c for c in line.split("|") if c != ""]
            # separator row (e.g. |---|---|)
            if all(re.match(r"[-: ]+$", c.strip()) for c in cols):
                i += 1
                continue
            if not in_table:
                in_table      = True
                table_headers = [_strip_inline(c) for c in cols]
            else:
                table_rows.append([_strip_inline(c) for c in cols])
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── headings ──────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = _strip_inline(m.group(2))
            if level == 1:
                pass   # already used as doc title in cover
            elif level == 2:
                _h2(doc, text)
            elif level == 3:
                _h3(doc, text)
            else:
                _h3(doc, text)
            i += 1
            continue

        # ── horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^[-*_]{3,}\s*$", line.strip()):
            _add_hr(doc)
            i += 1
            continue

        # ── bullet list ───────────────────────────────────────────────────────
        bm = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if bm:
            level = len(bm.group(1)) // 2
            _bullet(doc, _strip_inline(bm.group(2)), level=level)
            i += 1
            continue

        # ── numbered list ─────────────────────────────────────────────────────
        nm = re.match(r"^\s*\d+\.\s+(.*)", line)
        if nm:
            _bullet(doc, _strip_inline(nm.group(1)))
            i += 1
            continue

        # ── blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── plain / inline-styled paragraph ───────────────────────────────────
        _inline_styled(doc, line.strip())
        i += 1

    # flush any open state
    if in_table:
        flush_table()
    if in_code:
        flush_code()

    doc.save(str(out_path))


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    md_files = sorted(
        p for p in ROOT_DIR.rglob("*.md")
        if not any(part in EXCLUDE_DIRS for part in p.parts)
    )
    print(f"Found {len(md_files)} markdown file(s)\n")

    for md_path in md_files:
        out_name = md_path.stem + ".docx"
        out_path = OUTPUT_DIR / out_name
        try:
            md_to_docx(md_path, out_path)
            print(f"  [OK]  {md_path.relative_to(ROOT_DIR)}  →  presentation_file/{out_name}")
        except Exception as exc:
            print(f"  [ERR] {md_path.relative_to(ROOT_DIR)}: {exc}")

    print(f"\nAll files saved to: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
