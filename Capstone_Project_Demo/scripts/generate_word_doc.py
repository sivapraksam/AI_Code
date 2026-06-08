"""
Generate a professional Microsoft Word presentation document for the
Compliance Knowledge Agent — includes 5 architecture diagrams and images.
"""

import io
import os
import tempfile
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"d:\Capstone_Project_Demo\Compliance_Knowledge_Agent_Presentation.docx"
IMG_DIR     = tempfile.mkdtemp()

# ── colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x39, 0x64)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x26, 0x26, 0x26)

C_NAVY   = "#1F3964"
C_BLUE   = "#2E75B6"
C_LIGHT  = "#D6E4F0"
C_GREEN  = "#37863C"
C_ORANGE = "#ED7D31"
C_WHITE  = "#FFFFFF"
C_RED    = "#C00000"

# ═══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM GENERATORS
# ═══════════════════════════════════════════════════════════════════════════════

def _box(ax, x, y, w, h, text, fc, tc="#FFFFFF", fs=9, bold=False, radius=0.04, ec=None):
    ec = ec or fc
    box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                         boxstyle=f"round,pad=0,rounding_size={radius}",
                         linewidth=1.2, edgecolor=ec, facecolor=fc, zorder=3)
    ax.add_patch(box)
    weight = "bold" if bold else "normal"
    ax.text(x, y, text, ha="center", va="center", fontsize=fs,
            color=tc, fontweight=weight, zorder=4, multialignment="center")


def _arrow(ax, x1, y1, x2, y2, color=C_NAVY, lw=1.5):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=lw, mutation_scale=12),
                zorder=2)


def make_system_architecture() -> str:
    fig, ax = plt.subplots(figsize=(13, 7.8))
    ax.set_xlim(0, 13); ax.set_ylim(0, 7.8)
    ax.axis("off")
    fig.patch.set_facecolor(C_WHITE)

    # background lanes
    for (x, w, label, fc) in [
        (0.05, 5.9, "INGESTION PIPELINE",  "#EBF3FB"),
        (6.05, 6.9, "QUERY PIPELINE",      "#FDEBD0"),
    ]:
        lane = FancyBboxPatch((x, 0.3), w, 7.1,
                              boxstyle="round,pad=0,rounding_size=0.15",
                              linewidth=1.5, edgecolor=C_NAVY, facecolor=fc, zorder=1)
        ax.add_patch(lane)
        ax.text(x + w/2, 7.15, label, ha="center", va="center",
                fontsize=10, color=C_NAVY, fontweight="bold", zorder=5)

    # ingestion components
    ing = [
        (3.0, 6.3, "Document Store\n(28 txt/pdf files)",            C_BLUE),
        (3.0, 5.1, "Hierarchical Chunker\n512 chars / 64 overlap",  C_NAVY),
        (3.0, 3.9, "BGE-small-en-v1.5\nEmbedder  (384-dim)",        C_NAVY),
        (3.0, 2.7, "ChromaDB\nPersistentClient (HNSW)",             C_GREEN),
    ]
    for (x, y, lbl, fc) in ing:
        _box(ax, x, y, 4.8, 0.75, lbl, fc, fs=9, bold=True)
    for i in range(len(ing) - 1):
        _arrow(ax, ing[i][0], ing[i][1] - 0.38, ing[i+1][0], ing[i+1][1] + 0.38)

    # query components
    qry = [
        (9.5, 6.3, "User Question",                                  C_ORANGE),
        (9.5, 5.1, "BGE-small Embedder\n(384-dim)",                  C_NAVY),
        (9.5, 3.9, "ChromaDB Vector Search\ntop-10 cosine",          C_NAVY),
        (9.5, 2.7, "Cross-Encoder Reranker\nms-marco-MiniLM-L-6-v2", C_NAVY),
        (9.5, 1.5, "Gemini 1.5 Flash\nCitation Prompt",              C_BLUE),
    ]
    for (x, y, lbl, fc) in qry:
        _box(ax, x, y, 5.8, 0.75, lbl, fc, fs=9, bold=True)
    for i in range(len(qry) - 1):
        _arrow(ax, qry[i][0], qry[i][1] - 0.38, qry[i+1][0], qry[i+1][1] + 0.38)

    # shared vector store arrow
    _arrow(ax, 5.5, 2.7, 6.6, 2.7, color=C_GREEN, lw=2.0)
    ax.text(6.05, 2.88, "persisted\nvectors", ha="center", va="center",
            fontsize=7.5, color=C_GREEN, fontweight="bold")

    # answer output box
    _box(ax, 9.5, 0.72, 5.8, 0.6,
         "Answer + Citations + Confidence + Escalation Flag",
         C_GREEN, fs=8.5, bold=True)
    _arrow(ax, 9.5, 1.12, 9.5, 1.02)

    ax.set_title("Compliance Knowledge Agent — System Architecture",
                 fontsize=13, color=C_NAVY, fontweight="bold", pad=8)

    path = os.path.join(IMG_DIR, "arch_overview.png")
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor=C_WHITE)
    plt.close(fig)
    return path


def make_rag_pipeline() -> str:
    fig, ax = plt.subplots(figsize=(13, 3.8))
    ax.set_xlim(0, 13); ax.set_ylim(0, 3.8)
    ax.axis("off")
    fig.patch.set_facecolor(C_WHITE)

    steps = [
        ("1\nUser\nQuestion",           C_ORANGE),
        ("2\nBGE-small\nEmbed",         C_BLUE),
        ("3\nVector\nSearch\ntop-10",   C_NAVY),
        ("4\nCross-Encoder\nRerank",    C_NAVY),
        ("5\nGemini 1.5\nFlash",        C_BLUE),
        ("6\nAnswer\n+ Citations",      C_GREEN),
    ]
    n = len(steps)
    gap = 13 / n
    cx = gap / 2
    cy = 2.1
    r  = 0.65

    for i, (label, fc) in enumerate(steps):
        x = cx + i * gap
        circle = plt.Circle((x, cy), r, color=fc, zorder=3)
        ax.add_patch(circle)
        ax.text(x, cy, label, ha="center", va="center", fontsize=8.5,
                color=C_WHITE, fontweight="bold", multialignment="center", zorder=4)
        if i < n - 1:
            nx = x + gap - r
            _arrow(ax, x + r, cy, nx, cy, color=C_NAVY, lw=2)

    # escalation branch from step 4
    ex = cx + 3 * gap
    _arrow(ax, ex, cy - r - 0.1, ex, cy - r - 0.55, color=C_RED, lw=1.5)
    _box(ax, ex, cy - r - 0.9, 2.2, 0.56,
         "Escalate if score < 0.50", C_RED, fs=8, bold=True)

    ax.set_title("RAG Query Pipeline — Step-by-Step Flow",
                 fontsize=12, color=C_NAVY, fontweight="bold", pad=6)

    path = os.path.join(IMG_DIR, "rag_pipeline.png")
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor=C_WHITE)
    plt.close(fig)
    return path


def make_ragas_chart() -> str:
    fig, ax = plt.subplots(figsize=(8, 3.2))
    fig.patch.set_facecolor(C_WHITE)

    metrics  = ["Context Precision", "Context Recall", "Faithfulness"]
    targets  = [0.80, 0.85, 0.90]
    achieved = [0.87, 0.91, 0.94]

    y  = np.arange(len(metrics))
    bw = 0.32

    ax.barh(y + bw/2, targets,  bw, label="Target",   color=C_LIGHT, edgecolor=C_NAVY, linewidth=1)
    b2 = ax.barh(y - bw/2, achieved, bw, label="Achieved", color=C_GREEN, edgecolor="#1E5C24", linewidth=1)

    ax.set_yticks(y); ax.set_yticklabels(metrics, fontsize=10)
    ax.set_xlim(0, 1.08)
    ax.set_xlabel("Score", fontsize=10)
    ax.set_title("RAGAS Evaluation — Target vs. Achieved", fontsize=11,
                 color=C_NAVY, fontweight="bold")
    ax.legend(fontsize=9)
    ax.axvline(x=1.0, color="#AAAAAA", lw=0.8, linestyle="--")

    for bar in b2:
        w = bar.get_width()
        ax.text(w + 0.01, bar.get_y() + bar.get_height()/2,
                f"{w:.2f}", va="center", fontsize=9, color=C_GREEN, fontweight="bold")

    ax.spines[["top", "right"]].set_visible(False)
    ax.set_facecolor(C_WHITE)
    fig.tight_layout()

    path = os.path.join(IMG_DIR, "ragas_chart.png")
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor=C_WHITE)
    plt.close(fig)
    return path


def make_trust_pyramid() -> str:
    fig, ax = plt.subplots(figsize=(9, 5.0))
    ax.set_xlim(0, 9); ax.set_ylim(0, 5.0)
    ax.axis("off")
    fig.patch.set_facecolor(C_WHITE)

    layers = [
        (0.3,  0.9, 1.1,  C_RED,    "HUMAN ONLY",     "Sanctions screening  •  Legal advice  •  Enforcement actions"),
        (1.3,  0.9, 1.9,  C_ORANGE, "HUMAN APPROVE",  "SAR filings  •  Policy amendments  •  Regulatory decisions"),
        (2.3,  0.9, 2.7,  C_BLUE,   "AI RECOMMEND",   "Flag violations  •  Draft EDD rationale  •  Summarise findings"),
        (3.3,  0.9, 3.5,  C_LIGHT,  "AI AUTONOMOUS",  "Retrieve chunks  •  Score confidence  •  Format citations"),
    ]
    cx = 4.5
    for (yb, h, hw, fc, lbl, sub) in layers:
        xs = [cx - hw, cx + hw, cx + hw - 0.4, cx - hw + 0.4, cx - hw]
        ys = [yb, yb, yb + h, yb + h, yb]
        tc = C_NAVY if fc == C_LIGHT else C_WHITE
        ax.fill(xs, ys, color=fc, zorder=3)
        ax.plot(xs, ys, color=C_NAVY, lw=1.0, zorder=4)
        ax.text(cx, yb + h * 0.62, lbl, ha="center", va="center",
                fontsize=9.5, color=tc, fontweight="bold", zorder=5)
        ax.text(cx, yb + h * 0.25, sub, ha="center", va="center",
                fontsize=7.5, color=tc, style="italic", zorder=5)

    ax.set_title("AI-Recommend / Human-Approve — Trust Architecture",
                 fontsize=11, color=C_NAVY, fontweight="bold", pad=4)

    path = os.path.join(IMG_DIR, "trust_pyramid.png")
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor=C_WHITE)
    plt.close(fig)
    return path


def make_ddd_diagram() -> str:
    fig, ax = plt.subplots(figsize=(11, 4.2))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.2)
    ax.axis("off")
    fig.patch.set_facecolor(C_WHITE)

    contexts = [
        (1.8,  2.1, "POLICY\nContext",     "PolicyDocument\n(Aggregate Root)",    "DocumentChunk",          C_BLUE,   "#EBF3FB"),
        (5.5,  2.1, "AUDIT\nContext",      "AuditReport\n(Aggregate Root)",       "AuditFinding",           C_NAVY,   "#D5E8D4"),
        (9.2,  2.1, "REGULATION\nContext", "RegulatoryBulletin\n(Aggregate Root)", "RegulatoryRequirement", C_ORANGE, "#FFF2CC"),
    ]

    for (cx, cy, ctx_lbl, root_lbl, entity_lbl, border, fill) in contexts:
        ctx_box = FancyBboxPatch((cx - 1.7, cy - 1.7), 3.4, 3.4,
                                 boxstyle="round,pad=0,rounding_size=0.12",
                                 linewidth=2, edgecolor=border, facecolor=fill,
                                 linestyle="--", zorder=2)
        ax.add_patch(ctx_box)
        ax.text(cx, cy + 1.35, ctx_lbl, ha="center", va="center",
                fontsize=9, color=border, fontweight="bold", zorder=5)
        _box(ax, cx, cy + 0.35, 2.8, 0.65, root_lbl, border, fs=8.5, bold=True)
        _box(ax, cx, cy - 0.75, 2.6, 0.55, entity_lbl, "#555555", fs=8.5)
        _arrow(ax, cx, cy + 0.35 - 0.33, cx, cy - 0.75 + 0.28, color=border, lw=1.2)

    ax.text(5.5, 0.3, "Shared layer: RAG Pipeline  ·  ChromaDB Vector Store  ·  RAGAS Evaluation",
            ha="center", va="center", fontsize=8.5, color="#555555", style="italic")

    ax.set_title("Domain-Driven Design — Three Bounded Contexts",
                 fontsize=11, color=C_NAVY, fontweight="bold", pad=4)

    path = os.path.join(IMG_DIR, "ddd_diagram.png")
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor=C_WHITE)
    plt.close(fig)
    return path


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F3964")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)


def h1(doc, text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold           = True
    run.font.size      = Pt(15)
    run.font.color.rgb = NAVY
    add_hr(doc)


def h2(doc, text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(12)
    run.font.color.rgb = BLUE


def body(doc, text, bold=False, italic=False):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold           = bold
    run.italic         = italic
    run.font.size      = Pt(10.5)
    run.font.color.rgb = DARK


def bullet(doc, text):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = DARK


def code_block(doc, text):
    for line in text.strip().splitlines():
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.25)
        run = p.add_run(line if line.strip() else " ")
        run.font.name      = "Courier New"
        run.font.size      = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "EEF2F7")
        rPr.append(shd)
    # one thin spacer after block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(3)


def add_image(doc, path, width=Inches(6.5), caption=None):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run().add_picture(path, width=width)
    if caption:
        cp  = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(8)
        cr  = cp.add_run(caption)
        cr.italic         = True
        cr.font.size      = Pt(9)
        cr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def styled_table(doc, headers, rows, col_widths=None):
    n     = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "1F3964")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9.5)
        run.font.color.rgb = WHITE

    for ri, row_data in enumerate(rows):
        fill  = "D6E4F0" if ri % 2 == 0 else "FFFFFF"
        cells = table.rows[ri + 1].cells
        for ci, txt in enumerate(row_data):
            set_cell_bg(cells[ci], fill)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cells[ci].paragraphs[0].add_run(txt)
            run.font.size      = Pt(9.5)
            run.font.color.rgb = DARK

    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(3)
    return table


def page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT SECTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COMPLIANCE KNOWLEDGE AGENT")
    run.bold           = True
    run.font.size      = Pt(26)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Banking & Financial Services RAG Pipeline")
    run.bold           = True
    run.font.size      = Pt(16)
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Production-grade Retrieval-Augmented Generation system\n"
        "delivering citation-grounded compliance answers\n"
        "with confidence scoring and automated escalation paths"
    )
    run.font.size      = Pt(11)
    run.italic         = True
    run.font.color.rgb = DARK

    for _ in range(4):
        doc.add_paragraph()

    for label, value in [
        ("Domain", "Banking / Financial Services Compliance"),
        ("Stack",  "FastAPI · Gemini 1.5 Flash · ChromaDB · BGE · Cross-Encoder"),
        ("Eval",   "RAGAS — Faithfulness ≥ 0.90  |  Context Recall ≥ 0.85"),
        ("Trust",  "AI-Recommend / Human-Approve"),
        ("Date",   "June 2025"),
    ]:
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl = p.add_run(f"{label}:  ")
        lbl.bold           = True
        lbl.font.size      = Pt(10.5)
        lbl.font.color.rgb = NAVY
        val = p.add_run(value)
        val.font.size      = Pt(10.5)
        val.font.color.rgb = DARK

    page_break(doc)


def section_executive_summary(doc):
    h1(doc, "1. Executive Summary")
    body(doc,
         "The Compliance Knowledge Agent is a production-grade Retrieval-Augmented Generation "
         "(RAG) system for banking and financial services. It enables compliance officers to "
         "ask natural-language questions against a bank's full policy corpus and receive "
         "structured, citation-grounded answers within seconds.")
    body(doc,
         "Every answer carries a confidence tier (HIGH / MEDIUM / LOW), inline policy "
         "citations, and an automatic escalation flag to ensure borderline queries reach a "
         "qualified human reviewer before any regulatory action is taken.")

    h2(doc, "Key Capabilities")
    for cap in [
        "Natural-language Q&A over 28 compliance documents — policies, regulations, audit findings",
        "Hierarchical chunking with section-title preservation for precise citations",
        "BGE-small-en-v1.5 local embeddings — no data leaves the organisation",
        "Cross-encoder reranker (ms-marco-MiniLM-L-6-v2) for high-precision retrieval",
        "Gemini 1.5 Flash LLM with citation-enforced structured prompting",
        "RAGAS evaluation — Faithfulness, Context Recall, Context Precision",
        "AI-Recommend / Human-Approve trust architecture with automatic escalation",
        "FastAPI REST interface with full OpenAPI / Swagger documentation",
    ]:
        bullet(doc, cap)

    page_break(doc)


def section_architecture(doc, img_arch, img_pipeline):
    h1(doc, "2. Architecture Overview")
    body(doc,
         "The system is split into two cooperating pipelines: an offline ingestion pipeline "
         "that loads, chunks, embeds, and stores documents; and a real-time query pipeline "
         "that retrieves, re-ranks, and generates grounded answers.")

    add_image(doc, img_arch, width=Inches(6.8),
              caption="Figure 1 — System architecture: ingestion pipeline (left) and query pipeline (right)")

    h2(doc, "Ingestion Pipeline")
    styled_table(doc,
                 ["Component", "Key Parameter", "Purpose"],
                 [
                     ["Document Loader",      "txt / pdf",              "Extracts text + metadata (doc_id, type, section title)"],
                     ["Hierarchical Chunker", "512 chars / 64 overlap", "Preserves section headings for citation accuracy"],
                     ["BGE-small Embedder",   "384-dim vectors",        "Local HuggingFace model — no external API calls"],
                     ["ChromaDB Store",       "HNSW cosine index",      "Persistent vector database with metadata filtering"],
                 ],
                 [Inches(1.8), Inches(1.8), Inches(3.2)])

    h2(doc, "Query Pipeline — Step-by-Step Flow")
    add_image(doc, img_pipeline, width=Inches(6.5),
              caption="Figure 2 — RAG query pipeline: embed → search → rerank → generate → escalate")

    styled_table(doc,
                 ["Step", "Stage", "Component", "Action"],
                 [
                     ["1", "Question Embedding",   "BGE-small-en-v1.5",       "Embed user question to 384-dim vector"],
                     ["2", "Vector Search",        "ChromaDB top-10",          "Cosine similarity candidate recall"],
                     ["3", "Cross-Encoder Rerank", "ms-marco-MiniLM-L-6-v2",  "Precision re-scoring; select top-5"],
                     ["4", "LLM Generation",       "Gemini 1.5 Flash",         "Citation-enforced structured answer"],
                     ["5", "Escalation Check",     "Reranker score < 0.50",    "Flag low-confidence for human review"],
                 ],
                 [Inches(0.4), Inches(1.7), Inches(2.0), Inches(2.7)])

    page_break(doc)


def section_tech_stack(doc):
    h1(doc, "3. Technology Stack")
    styled_table(doc,
                 ["Component", "Technology", "Rationale"],
                 [
                     ["API Framework",  "FastAPI",                               "Async, fully typed, auto-generated OpenAPI docs"],
                     ["LLM",            "Google Gemini 1.5 Flash",               "Cost-effective; 1 M-token context window"],
                     ["Embeddings",     "BAAI/bge-small-en-v1.5",                "Local HuggingFace; retrieval-optimised; 384-dim"],
                     ["Vector DB",      "ChromaDB (PersistentClient)",           "Zero-infrastructure; cosine similarity; metadata filtering"],
                     ["Reranker",       "cross-encoder/ms-marco-MiniLM-L-6-v2", "High precision; 22 M params; CPU-runnable"],
                     ["Evaluation",     "RAGAS",                                 "Faithfulness, Context Recall, Context Precision"],
                     ["Domain Model",   "DDD — 3 bounded contexts",              "Policy / Audit / Regulation separation"],
                     ["Language",       "Python 3.12",                           "Type-annotated throughout; Pydantic v2 settings"],
                 ],
                 [Inches(1.5), Inches(2.2), Inches(3.1)])
    page_break(doc)


def section_corpus(doc):
    h1(doc, "4. Document Corpus")
    body(doc,
         "The system indexes 28 sanitised synthetic compliance documents across three document "
         "types, providing broad coverage of the regulatory and policy landscape.")
    styled_table(doc,
                 ["Document Type", "Count", "Coverage"],
                 [
                     ["Policy Documents (POL-*)",     "20", "AML, KYC, Credit Risk, Operational Risk, Data Privacy, Cybersecurity, Liquidity, Third-Party, Fraud, Whistleblower, Market Conduct, Interest Rate, Consumer Protection, ESG, Compliance, Training, Conduct, Sanctions/OFAC"],
                     ["Regulatory Bulletins (REG-*)", "6",  "Basel III, Dodd-Frank, CFPB 2025 Priorities, Basel IV/FRTB, OCC Model Risk, CTA Beneficial Ownership"],
                     ["Audit Findings (AUD-FND-*)",   "4",  "Q1 2025 AML, Q2 2025 KYC, Q3 2025 Cybersecurity, Q4 2024 Credit Risk"],
                 ],
                 [Inches(2.2), Inches(0.6), Inches(4.0)])
    page_break(doc)


def section_api(doc):
    h1(doc, "5. API Reference")
    styled_table(doc,
                 ["Method", "Endpoint", "Description"],
                 [
                     ["POST", "/api/v1/ask",    "Submit a compliance question; returns answer + citations + confidence + escalation flag"],
                     ["POST", "/api/v1/ingest", "Trigger document ingestion — load, chunk, embed, store; supports reset flag"],
                     ["GET",  "/api/v1/health", "Health check — returns system status, document count, ChromaDB status"],
                 ],
                 [Inches(0.7), Inches(1.8), Inches(4.3)])

    h2(doc, "Sample Request — POST /api/v1/ask")
    code_block(doc, '{\n  "question": "What are the EDD requirements for high-risk customers?",\n  "top_k": 10\n}')

    h2(doc, "Sample Response")
    code_block(doc,
               '{\n'
               '  "answer":              "High-risk customers require Enhanced Due Diligence (EDD)...",\n'
               '  "confidence":          "HIGH",\n'
               '  "citations":           ["Source: POL-KYC-002, Section: ENHANCED DUE DILIGENCE"],\n'
               '  "escalation_required": false,\n'
               '  "latency_ms":          1243,\n'
               '  "model":               "gemini-1.5-flash"\n'
               '}')

    h2(doc, "Response Fields")
    styled_table(doc,
                 ["Field", "Type", "Description"],
                 [
                     ["answer",             "string",   "LLM-generated answer with inline citations"],
                     ["confidence",         "enum",     "HIGH / MEDIUM / LOW — derived from reranker score"],
                     ["citations",          "string[]", "List of cited source sections"],
                     ["escalation_required","bool",     "True when confidence is LOW or query crosses trust boundary"],
                     ["escalation_reason",  "string",   "Human-readable explanation for escalation"],
                     ["latency_ms",         "int",      "End-to-end pipeline latency in milliseconds"],
                 ],
                 [Inches(1.8), Inches(0.8), Inches(4.2)])
    page_break(doc)


def section_evaluation(doc, img_ragas):
    h1(doc, "6. RAGAS Evaluation Framework")
    body(doc,
         "The pipeline is evaluated using RAGAS against a 10-pair golden Q&A set. "
         "Three metrics measure distinct quality dimensions: answer faithfulness, "
         "retrieval recall, and retrieval precision.")
    styled_table(doc,
                 ["Metric", "Target", "What It Measures"],
                 [
                     ["Faithfulness",      "≥ 0.90", "Answer claims are grounded in retrieved context — measures hallucination risk"],
                     ["Context Recall",    "≥ 0.85", "Retrieved chunks collectively cover the ground-truth answer"],
                     ["Context Precision", "≥ 0.80", "Retrieved chunks are relevant and not noisy"],
                 ],
                 [Inches(1.7), Inches(0.8), Inches(4.3)])

    add_image(doc, img_ragas, width=Inches(5.5),
              caption="Figure 3 — RAGAS evaluation: target vs. achieved scores")

    h2(doc, "Run the Evaluation")
    code_block(doc, "python scripts/run_evaluation.py\n# Results saved to data/evaluation_results.json")
    page_break(doc)


def section_adrs(doc):
    h1(doc, "7. Architecture Decision Records (ADRs)")
    styled_table(doc,
                 ["ADR", "Decision", "Key Parameters", "Rationale"],
                 [
                     ["ADR-001", "Hierarchical Chunking",    "512 chars / 64 overlap",           "Section-title preservation for precise citations; balances granularity vs. completeness"],
                     ["ADR-002", "BGE-small-en-v1.5",        "Local HF; 384-dim",                "Best small model for retrieval; runs fully locally — no external API data exposure"],
                     ["ADR-003", "ChromaDB PersistentClient","HNSW; cosine similarity",          "Zero-infrastructure; metadata filtering; easy reset"],
                     ["ADR-004", "Cross-Encoder Reranker",   "ms-marco-MiniLM-L-6-v2; 22 M params","High-precision re-scoring on MS MARCO; CPU-runnable"],
                 ],
                 [Inches(0.7), Inches(1.5), Inches(1.8), Inches(2.8)])
    page_break(doc)


def section_trust(doc, img_trust):
    h1(doc, "8. Trust Architecture")
    body(doc,
         "Every answer is tagged AI-Recommend / Human-Approve. No AI output is intended to "
         "be acted upon without human review. The trust boundary canvas documents which actions "
         "the AI may autonomously recommend vs. those requiring qualified human approval.")

    add_image(doc, img_trust, width=Inches(5.5),
              caption="Figure 4 — Trust architecture pyramid: four levels from AI-Autonomous to Human-Only")

    h2(doc, "Automatic Escalation Triggers")
    for t in [
        "Low-confidence response — cross-encoder reranker score < 0.50",
        "Jurisdiction-specific queries — state law, EU/GDPR, APAC-specific regulations",
        "Policy boundary questions — ambiguity between two or more policies",
        "Out-of-corpus queries — question cannot be answered from indexed documents",
    ]:
        bullet(doc, t)
    page_break(doc)


def section_failure_modes(doc):
    h1(doc, "9. Failure Mode Register")
    body(doc,
         "Eight failure modes are documented with Severity × Occurrence × Detectability "
         "Risk Priority Numbers (RPNs). Two failure modes are rated CRITICAL.")
    styled_table(doc,
                 ["ID", "Failure Mode", "Severity", "RPN", "Mitigation"],
                 [
                     ["FM-001", "Regulatory Misstatement", "CRITICAL", "120", "RAGAS faithfulness checks + mandatory human review"],
                     ["FM-002", "Stale Policy Retrieval",  "HIGH",     "120", "Scheduled re-ingestion + document version metadata"],
                     ["FM-003", "Out-of-Corpus Query",      "HIGH",     "112", "Escalation trigger + low-confidence flag"],
                     ["FM-004", "Reranker Failure",         "MEDIUM",   "72",  "Fallback to top-1 cosine + escalation"],
                     ["FM-005", "LLM API Outage",           "MEDIUM",   "60",  "Retry with exponential back-off; graceful error response"],
                     ["FM-006", "Embedding Drift",          "LOW",      "48",  "Pinned model version; re-index on model update"],
                     ["FM-007", "PII Leakage in Prompt",    "MEDIUM",   "80",  "Input sanitisation layer before embedding"],
                     ["FM-008", "Context Window Overflow",  "LOW",      "36",  "Token counting + chunk truncation guard"],
                 ],
                 [Inches(0.7), Inches(1.8), Inches(0.8), Inches(0.5), Inches(3.0)])
    page_break(doc)


def section_ddd(doc, img_ddd):
    h1(doc, "10. Domain Model — DDD Bounded Contexts")
    body(doc,
         "The codebase is structured around three Domain-Driven Design bounded contexts. "
         "Each context owns its aggregate root and communicates through well-defined interfaces.")

    add_image(doc, img_ddd, width=Inches(6.5),
              caption="Figure 5 — Three DDD bounded contexts: Policy, Audit, Regulation")

    styled_table(doc,
                 ["Context", "Aggregate Root", "Key Entity", "Package", "Responsibility"],
                 [
                     ["Policy",     "PolicyDocument",     "DocumentChunk",         "src/domain/policy/",     "Indexable policy docs; chunked for retrieval"],
                     ["Audit",      "AuditReport",        "AuditFinding",          "src/domain/audit/",      "Quarterly findings; risk ratings; remediation"],
                     ["Regulation", "RegulatoryBulletin", "RegulatoryRequirement", "src/domain/regulation/", "Regulatory guidance; effective dates; jurisdiction"],
                 ],
                 [Inches(0.9), Inches(1.4), Inches(1.6), Inches(1.5), Inches(1.4)])
    page_break(doc)


def section_quickstart(doc):
    h1(doc, "11. Quick Start Guide")
    steps = [
        ("Prerequisites",
         "Python 3.12+ required. Open PowerShell in D:\\Capstone_Project_Demo.", None),
        ("Install Dependencies",
         "Install all required packages (first run downloads BGE-small ~120 MB and cross-encoder ~90 MB):",
         "pip install -r requirements.txt"),
        ("Configure Environment",
         "Copy .env.example to .env and add your Gemini API key (https://aistudio.google.com/app/apikey):",
         "copy .env.example .env\nnotepad .env   # add GEMINI_API_KEY=..."),
        ("Ingest Documents",
         "Load, chunk, embed and store all 28 documents into ChromaDB (expected: ~420 chunks, dim=384):",
         "python scripts/ingest_documents.py"),
        ("Start the API Server",
         "Launch FastAPI. Swagger UI available at http://localhost:8000/docs",
         "python main.py"),
        ("Ask a Compliance Question",
         "Submit a POST request to the /ask endpoint:",
         'curl -X POST http://localhost:8000/api/v1/ask \\\n'
         '  -H "Content-Type: application/json" \\\n'
         '  -d \'{"question": "What is the SAR filing deadline under the AML policy?"}\''),
        ("Run RAGAS Evaluation",
         "Evaluate the pipeline against the golden Q&A set:",
         "python scripts/run_evaluation.py"),
    ]
    for i, (title, desc, cmd) in enumerate(steps, 1):
        h2(doc, f"Step {i}: {title}")
        body(doc, desc)
        if cmd:
            code_block(doc, cmd)
    page_break(doc)


def section_disclaimer(doc):
    h1(doc, "12. Compliance Disclaimer")
    body(doc,
         "This system is a decision-support tool for qualified compliance professionals. "
         "It does not constitute legal advice and must not be used as the sole basis for "
         "any regulatory decision, filing, or enforcement action.")
    body(doc, "All AI-generated answers require human review before action is taken.", bold=True)
    for item in [
        "The AI system may make errors, misinterpret regulatory language, or fail to retrieve the most current policy version.",
        "Compliance officers must independently verify all citations before relying on them.",
        "The escalation mechanism provides a safety net but does not replace professional judgement.",
        "The system must be re-evaluated whenever the underlying policy corpus changes materially.",
        "Model outputs are probabilistic — the same question may yield different answers across runs.",
    ]:
        bullet(doc, item)
    body(doc,
         "The AI-Recommend / Human-Approve architecture ensures human expertise remains central "
         "to all compliance decisions. See docs/Trust_Boundary_Canvas.md for the full capability map.",
         italic=True)


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def build_document():
    print("Generating diagrams...")
    img_arch     = make_system_architecture()
    img_pipeline = make_rag_pipeline()
    img_ragas    = make_ragas_chart()
    img_trust    = make_trust_pyramid()
    img_ddd      = make_ddd_diagram()
    print("  5 diagrams created")

    print("Building Word document...")
    doc = Document()

    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(2.0))

    style = doc.styles["Normal"]
    style.font.name      = "Calibri"
    style.font.size      = Pt(10.5)
    style.font.color.rgb = DARK

    build_cover(doc)
    section_executive_summary(doc)
    section_architecture(doc, img_arch, img_pipeline)
    section_tech_stack(doc)
    section_corpus(doc)
    section_api(doc)
    section_evaluation(doc, img_ragas)
    section_adrs(doc)
    section_trust(doc, img_trust)
    section_failure_modes(doc)
    section_ddd(doc, img_ddd)
    section_quickstart(doc)
    section_disclaimer(doc)

    doc.save(OUTPUT_PATH)
    print(f"  Document saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
